import datetime
import logging
import traceback
import json
import re
import StringIO

import webapp2
from google.appengine.ext.webapp import template
import cloudstorage as gcs

import utils
from docx import Document
from .models import *

class FloodDamageHandler(webapp2.RequestHandler):
    def generate(self):
        user_email = utils.authenticate_user(self, self.request.url, ["dhirenvjti@gmail.com"])
        if not user_email:
            return

        if self.request.method == 'GET':
            page = utils.template("flood_damage.html", "Report/html")
            template_values = {}
            self.response.out.write(template.render(page, template_values))

        else:
            self.response.headers['Content-Type'] = "application/json"
            self.response.headers['Access-Control-Allow-Origin'] = '*'
            try:
                date_of_sending = datetime.datetime.strptime(self.request.get("date_of_sending", None), "%Y-%m-%d")
                date_of_reporting = datetime.datetime.strptime(self.request.get("date_of_reporting", None), "%Y-%m-%d")
                replace_info = {
                    "date_of_sending": date_of_sending.strftime("%d/%m/%Y"),
                    "date_of_reporting": date_of_reporting.strftime("%d/%m/%Y"),
                    "general_trend_of_rainfall_last_24": self.request.get("general_trend_of_rainfall_last_24", None),
                    "cumulative_total_rainfall": self.request.get("cumulative_total_rainfall", None),
                    "river_name_levels_observed": self.request.get("river_name_levels_observed", None),
                    "name_engineering_works": self.request.get("name_engineering_works", None),
                    "communication_disruption_details": self.request.get("communication_disruption_details", None),
                    "relief_measures_taken": self.request.get("relief_measures_taken", None),
                    "comments_press_news": self.request.get("comments_press_news", None),
                    "enclosed_map": self.request.get("enclosed_map", None),
                    "flood_damaged_stats_attached": self.request.get("flood_damaged_stats_attached", None),
                    "sr_no": self.request.get("sr_no", None),
                    "state": self.request.get("state", "Dadra Nagar Haveli"),
                    "area_affected": self.request.get("area_affected", None),
                    "population_affected": self.request.get("population_affected", None),
                    "damage_to_crops": self.request.get("damage_to_crops", None),
                    "damage_to_houses": self.request.get("damage_to_houses", None),
                    "cattles_lost": self.request.get("cattles_lost", None),
                    "human_lives_lost": self.request.get("human_lives_lost", None),
                    "damage_to_public_utilities": self.request.get("damage_to_public_utilities", None),
                    "total_damage": self.request.get("total_damage", None),
                    "user_name": self.request.get("user_name", None),
                    "user_designation": self.request.get("user_designation", None),
                    "user_contact": self.request.get("user_contact", None),
                    "user_fax": self.request.get("user_fax", None),
                }

                document = Document(utils.template('Report-InformationrelatedtoFlood.docx', 'templates'))
                for phrase, replacement in replace_info.iteritems():
                    utils.docx_replace_regex(document, re.compile(phrase), replacement)

                target_stream = StringIO.StringIO()
                document.save(target_stream)

                created_at_IST = datetime.datetime.now(utils.TimeZone(+5.5, False, 'IST'))
                key = 'Report - Information related to Flood Stamp:{}'.format(created_at_IST.strftime("%Y-%m-%dT%H:%M"))
                write_retry_params = gcs.RetryParams(backoff_factor=1.1)
                filename = '/dnh-dma.appspot.com/%s.docx' % key
                gcs_file = gcs.open(filename, 'w', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', options={'x-goog-acl': 'public-read'},
                                    retry_params=write_retry_params)
                gcs_file.write(target_stream.getvalue())
                gcs_file.close()
                html_link = "https://storage.googleapis.com" + filename
                FloodDamageReport().add(
                    created_at_IST=created_at_IST,
                    report_link = html_link
                )
                self.response.out.write(json.dumps({'success': True, 'error': None, 'response': html_link}))
            except Exception as e:
                self.response.out.write(json.dumps({'success': False, 'error': e.message, 'response': None}))
                logging.error(traceback.format_exc())

    def download(self):
        flood_report = FloodDamageReport.get_latest_entry()
        self.redirect(str(flood_report.report_link))

class FloodSituationHandler(webapp2.RequestHandler):
    def generate(self):
        user_email = utils.authenticate_user(self, self.request.url, ["dhirenvjti@gmail.com"])
        if not user_email:
            return

        if self.request.method == 'GET':
            page = utils.template("flood_situation.html", "Report/html")
            template_values = {"default_date_of_reporting": datetime.datetime.now(utils.TimeZone(+5.5, False, 'IST')).strftime("%Y-%m-%dT%H:%M"),
                               "default_date_of_sending": datetime.datetime.now(utils.TimeZone(+5.5, False, 'IST')).strftime("%Y-%m-%d")}
            self.response.out.write(template.render(page, template_values))

        else:
            self.response.headers['Content-Type'] = "application/json"
            self.response.headers['Access-Control-Allow-Origin'] = '*'
            try:
                date_of_sending = datetime.datetime.strptime(self.request.get("date_of_sending", None), "%Y-%m-%d")
                date_of_reporting = datetime.datetime.strptime(self.request.get("date_of_reporting", None), "%Y-%m-%dT%H:%M")
                replace_info = {
                    "{date_of_sending}": date_of_sending.strftime("%d/%m/%Y"),
                    "{date_of_reporting}": date_of_reporting.strftime("%d/%m/%Y"),
                    "{time_of_reporting}": date_of_reporting.strftime("%H:%M"),
                    "{rainfall}": self.request.get("rainfall", None),
                    "{rainfall_last_24hrs}": self.request.get("rainfall_last_24hrs", None),
                    "{districts_affected}": self.request.get("districts_affected", None),
                    "{districts_affected_24hrs}": self.request.get("districts_affected_24hrs", None),
                    "{count_village_affected}": self.request.get("count_village_affected", None),
                    "{count_village_affected_24hrs}": self.request.get("count_village_affected_24hrs", None),
                    "{population_affected}": self.request.get("population_affected", None),
                    "{population_affected_24hrs}": self.request.get("population_affected_24hrs", None),
                    "{human_lives_lost}": self.request.get("human_lives_lost", None),
                    "{human_lives_lost_24hrs}": self.request.get("human_lives_lost_24hrs", None),
                    "{count_missing}": self.request.get("count_missing", None),
                    "{count_missing_24hrs}": self.request.get("count_missing_24hrs", None),
                    "{count_injured}": self.request.get("count_injured", None),
                    "{count_injured_24hrs}": self.request.get("count_injured_24hrs", None),
                    "{houses_damaged_fully}": self.request.get("houses_damaged_fully", None),
                    "{houses_damaged_partially}": self.request.get("houses_damaged_partially", None),
                    "{houses_damaged_24hrs}": self.request.get("houses_damaged_24hrs", None),
                    "{animal_deaths}": self.request.get("animal_deaths", None),
                    "{animal_deaths_24hrs}": self.request.get("animal_deaths_24hrs", None),
                    "{count_persons_evacuated}": self.request.get("count_persons_evacuated", None),
                    "{count_persons_evacuated_24hrs}": self.request.get("count_persons_evacuated_24hrs", None),
                    "{count_relief_camps}": self.request.get("count_relief_camps", None),
                    "{count_relief_camps_24hrs}": self.request.get("count_relief_camps_24hrs", None),
                    "{inmates_in_relief_camp}": self.request.get("inmates_in_relief_camp", None),
                    "{inmates_in_relief_camp_24hrs}": self.request.get("inmates_in_relief_camp_24hrs", None),
                    "{relief_material_distributed}": self.request.get("relief_material_distributed", None),
                    "{relief_material_distributed_24hrs}": self.request.get("relief_material_distributed_24hrs", None),
                    "{total_crop_area_affected}": self.request.get("total_crop_area_affected", None),
                    "{total_crop_area_affected_24hrs}": self.request.get("total_crop_area_affected_24hrs", None),
                    "{infra_damage}": self.request.get("infra_damage", None),
                    "{infra_damage_24hrs}": self.request.get("infra_damage_24hrs", None),
                    "{NDRF}": self.request.get("NDRF", None),
                    "{air_navy_army}": self.request.get("air_navy_army", None),
                    "{other_govt_dept}": self.request.get("other_govt_dept", None),
                    "{SDRF}": self.request.get("SDRF", None),
                    "{state_police_fire}": self.request.get("state_police_fire", None),
                    "{boats}": self.request.get("boats", None),
                    "{user_name}": self.request.get("user_name", None),
                    "{user_designation}": self.request.get("user_designation", None),
                    "{user_contact}": self.request.get("user_contact", None),
                    "{user_fax}": self.request.get("user_fax", None),
                }

                document = Document(utils.template('ReportonFlood-Situation.docx', 'templates'))
                for phrase, replacement in replace_info.iteritems():
                    utils.docx_replace_regex(document, re.compile(phrase), replacement)
                target_stream = StringIO.StringIO()
                document.save(target_stream)

                created_at_IST = datetime.datetime.now(utils.TimeZone(+5.5, False, 'IST'))
                key = 'Report on Flood-Situation Stamp:{}'.format(created_at_IST.strftime("%Y-%m-%dT%H:%M"))
                write_retry_params = gcs.RetryParams(backoff_factor=1.1)
                filename = '/dnh-dma.appspot.com/%s.docx' % key
                gcs_file = gcs.open(filename, 'w', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', options={'x-goog-acl': 'public-read'},
                                    retry_params=write_retry_params)
                gcs_file.write(target_stream.getvalue())
                gcs_file.close()
                html_link = "https://storage.googleapis.com" + filename
                FloodSituationReport().add(
                    created_at_IST=created_at_IST,
                    report_link = html_link
                )
                self.response.out.write(json.dumps({'success': True, 'error': None, 'response': html_link}))
            except Exception as e:
                self.response.out.write(json.dumps({'success': False, 'error': e.message, 'response': None}))
                logging.error(traceback.format_exc())

    def download(self):
        flood_situation = FloodSituationReport.get_latest_entry()
        self.redirect(str(flood_situation.report_link))
        
class NDMAHandler(webapp2.RequestHandler):
    def generate(self):
        user_email = utils.authenticate_user(self, self.request.url, ["dhirenvjti@gmail.com"])
        if not user_email:
            return

        if self.request.method == 'GET':
            page = utils.template("ndma.html", "Report/html")
            template_values = {"default_date_of_reporting": datetime.datetime.now(utils.TimeZone(+5.5, False, 'IST')).strftime("%Y-%m-%dT%H:%M"),
                               "default_date_of_sending": datetime.datetime.now(utils.TimeZone(+5.5, False, 'IST')).strftime("%Y-%m-%d")}
            self.response.out.write(template.render(page, template_values))

        else:
            self.response.headers['Content-Type'] = "application/json"
            self.response.headers['Access-Control-Allow-Origin'] = '*'
            try:
                date = datetime.datetime.strptime(self.request.get("date", None), "%Y-%m-%d")
                replace_info = {
                    "{date}": date.strftime("%d/%m/%Y"),
                    "{districts_affected}": self.request.get("districts_affected", None),
                    "{deaths}": self.request.get("deaths", None),
                    "{injured}": self.request.get("injured", None),
                    "{missing}": self.request.get("missing", None),
                    "{houses_damaged_fully}": self.request.get("houses_damaged_fully", None),
                    "{houses_damaged_partially}": self.request.get("houses_damaged_partially", None),
                    "{livestock_affected_big}": self.request.get("livestock_affected_big", None),
                    "{livestock_affected_small}": self.request.get("livestock_affected_small", None),
                    "{livestock_affected_poultry}": self.request.get("livestock_affected_poultry", None),
                    "{relief_camp}": self.request.get("relief_camp", None),
                    "{deployment_rescue_team}": self.request.get("deployment_rescue_team", None),
                    "{remarks}": self.request.get("remarks", None),
                    "{user_name}": self.request.get("user_name", None),
                    "{user_designation}": self.request.get("user_designation", None),
                    "{user_contact}": self.request.get("user_contact", None),
                    "{user_fax}": self.request.get("user_fax", None),
                }

                document = Document(utils.template('Report-NationalDisasterManagementAuthorityNDMA.docx', 'templates'))
                for phrase, replacement in replace_info.iteritems():
                    utils.docx_replace_regex(document, re.compile(phrase), replacement)
                target_stream = StringIO.StringIO()
                document.save(target_stream)

                created_at_IST = datetime.datetime.now(utils.TimeZone(+5.5, False, 'IST'))
                key = 'Report - National Disaster Management Authority (NDMA) Stamp:{}'.format(created_at_IST.strftime("%Y-%m-%dT%H:%M"))
                write_retry_params = gcs.RetryParams(backoff_factor=1.1)
                filename = '/dnh-dma.appspot.com/%s.docx' % key
                gcs_file = gcs.open(filename, 'w', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', options={'x-goog-acl': 'public-read'},
                                    retry_params=write_retry_params)
                gcs_file.write(target_stream.getvalue())
                gcs_file.close()
                html_link = "https://storage.googleapis.com" + filename
                NDMAReport().add(
                    created_at_IST=created_at_IST,
                    report_link = html_link
                )
                self.response.out.write(json.dumps({'success': True, 'error': None, 'response': html_link}))
            except Exception as e:
                self.response.out.write(json.dumps({'success': False, 'error': e.message, 'response': None}))
                logging.error(traceback.format_exc())

    def download(self):
        ndma = NDMAReport.get_latest_entry()
        self.redirect(str(ndma.report_link))
